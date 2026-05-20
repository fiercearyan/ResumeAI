from io import BytesIO
from docx import Document


def parse_docx(data: bytes) -> dict:
    doc = Document(BytesIO(data))
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    # Capture table cells too (some resumes use tables for layout).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)
    return {"lines": lines, "format": "docx"}
